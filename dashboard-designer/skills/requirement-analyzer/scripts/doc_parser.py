"""
文档解析脚本
支持解析Word和Excel文档
"""

import sys
import json
import argparse
from pathlib import Path

# 添加共享脚本路径
sys.path.insert(0, str(Path(__file__).parent.parent.parent / 'shared' / 'scripts'))

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


def parse_word(file_path: str) -> dict:
    """
    解析Word文档

    Args:
        file_path: Word文档路径

    Returns:
        包含文档内容的字典
    """
    if not HAS_DOCX:
        return {"error": "python-docx未安装，请运行: pip install python-docx"}

    doc = Document(file_path)
    content = {
        "paragraphs": [],
        "tables": [],
        "headings": []
    }

    # 提取段落
    for para in doc.paragraphs:
        if para.text.strip():
            content["paragraphs"].append({
                "text": para.text,
                "style": para.style.name if para.style else None
            })
            # 识别标题
            if para.style and 'Heading' in para.style.name:
                content["headings"].append({
                    "level": para.style.name,
                    "text": para.text
                })

    # 提取表格
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    return content


def parse_excel(file_path: str) -> dict:
    """
    解析Excel文档

    Args:
        file_path: Excel文档路径

    Returns:
        包含所有Sheet内容的字典
    """
    if not HAS_OPENPYXL:
        return {"error": "openpyxl未安装，请运行: pip install openpyxl"}

    wb = load_workbook(file_path, data_only=True)
    content = {
        "sheets": {}
    }

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_data = []

        for row in sheet.iter_rows(values_only=True):
            # 过滤空行
            if any(cell is not None for cell in row):
                sheet_data.append([str(cell) if cell is not None else "" for cell in row])

        content["sheets"][sheet_name] = sheet_data

    return content


def parse_document(file_path: str) -> dict:
    """
    根据文件扩展名自动选择解析器

    Args:
        file_path: 文档路径

    Returns:
        解析后的内容字典
    """
    ext = Path(file_path).suffix.lower()

    if ext in ['.docx', '.doc']:
        return parse_word(file_path)
    elif ext in ['.xlsx', '.xls']:
        return parse_excel(file_path)
    else:
        return {"error": f"不支持的文件格式: {ext}"}


def main():
    parser = argparse.ArgumentParser(description='文档解析工具')
    parser.add_argument('file_path', help='要解析的文档路径')
    parser.add_argument('--output', '-o', help='输出文件路径（JSON格式）')
    parser.add_argument('--format', '-f', choices=['json', 'text'], default='json',
                        help='输出格式')

    args = parser.parse_args()

    result = parse_document(args.file_path)

    if args.format == 'text':
        # 文本格式输出
        if 'paragraphs' in result:
            print("=== 段落内容 ===")
            for p in result['paragraphs']:
                print(p['text'])
        if 'tables' in result:
            print("\n=== 表格内容 ===")
            for i, table in enumerate(result['tables']):
                print(f"\n表格 {i+1}:")
                for row in table:
                    print(" | ".join(row))
        if 'sheets' in result:
            for sheet_name, data in result['sheets'].items():
                print(f"\n=== Sheet: {sheet_name} ===")
                for row in data:
                    print(" | ".join(row))
    else:
        # JSON格式输出
        output = json.dumps(result, ensure_ascii=False, indent=2)
        if args.output:
            Path(args.output).write_text(output, encoding='utf-8')
            print(f"结果已保存到: {args.output}")
        else:
            print(output)


if __name__ == '__main__':
    main()
